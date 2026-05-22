import html2text
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from werkzeug.utils import secure_filename
from openpyxl.styles import Font
import os, re
import fitz
from app import client
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import json
from flask import current_app
from app import db


from app.models.ai import Clauses
from app.models.auditOrganization import AuditOrganization
from app.models.project_instance_models import ProjectClause, ProjectComplianceActivity, ProjectControlActivity



def clean_html_preserve_format(html_content):
    if not html_content:
        return ""
    h = html2text.HTML2Text()
    h.ignore_links = True  # Optional: keep or remove links
    h.body_width = 0  # Do not wrap lines
    return h.handle(html_content).strip()


def format_consolidated_data_for_table(consolidated_data):
    """
    Format consolidated data for table display, ensuring proper text wrapping
    and handling of long content.
    """
    formatted_data = {}

    for key, value in consolidated_data.items():
        if value and value != "N/A":
            # Truncate very long content but preserve structure
            if len(str(value)) > 2000:
                # For very long content, truncate but keep important parts
                lines = str(value).split("\n")
                if len(lines) > 20:
                    # Keep first 15 lines and last 5 lines
                    truncated_lines = (
                        lines[:15] + ["...", "(content truncated)"] + lines[-5:]
                    )
                    formatted_data[key] = "\n".join(truncated_lines)
                else:
                    formatted_data[key] = str(value)[:1997] + "..."
            else:
                formatted_data[key] = str(value)
        else:
            formatted_data[key] = "N/A"

    return formatted_data


def add_consolidated_observation_table(document, observation_data):
    """
    Add a table for consolidated observation with improved formatting for consolidated data
    """
    # Add spacing before table
    document.add_paragraph()

    # Define fields in desired order
    field_order = [
        "Guideline Name",
        "Reference no.",
        "Clause Description",
        "Applicability Status",      
        "Assessment Status", 
        "Evidence Available", 
        "Test to be performed",
        "Detailed observation",
        "Gap noted & Impact",
        "Recommendation",
        "Compliance Status",
        "Severity",    
        "References to evidences/ Proof of concept",
    ]

    # Create table
    table = document.add_table(rows=len(field_order), cols=2)
    table.autofit = False
    table.allow_autofit = False

    # Set column widths
    table.columns[0].width = Inches(2.0)  # Field names
    table.columns[1].width = Inches(4.8)  # Content

    # Add borders
    add_table_borders(table)

    # Format and populate table
    formatted_data = format_consolidated_data_for_table(observation_data)

    for i, field_name in enumerate(field_order):
        row_cells = table.rows[i].cells

        # Field name cell
        field_paragraph = row_cells[0].paragraphs[0]
        field_run = field_paragraph.add_run(field_name)
        field_run.font.bold = True
        field_run.font.size = Pt(10)
        field_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Content cell
        content = formatted_data.get(field_name, "N/A")
        content_paragraph = row_cells[1].paragraphs[0]
        content_run = content_paragraph.add_run(content)
        content_run.font.size = Pt(9)
        content_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        content_paragraph.paragraph_format.left_indent = Inches(0.05)
        content_paragraph.paragraph_format.space_after = Pt(0)
        content_paragraph.paragraph_format.space_before = Pt(0)
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    print(
        f"Added consolidated table for clause: {observation_data.get('Reference no.', 'N/A')}"
    )



def add_severity_with_color(paragraph, severity):
    """Add severity text with appropriate color"""
    severity_colors = {
        'Critical': RGBColor(220, 38, 38),      # Red
        'Major': RGBColor(249, 115, 22),        # Orange
        'Significant': RGBColor(245, 158, 11),  # Yellow
        'Minor': RGBColor(96, 165, 250),        # Blue
        'No findings noted': RGBColor(34, 197, 94),  # Green
        'No Activities': RGBColor(156, 163, 175)     # Gray
    }
    
    run = paragraph.add_run(severity)
    if severity in severity_colors:
        run.font.color.rgb = severity_colors[severity]
    return run

def add_table_borders(table):
    """Add borders to all cells in a table with proper XML namespace"""
    # Add table borders using proper XML namespace
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    # Create table borders with proper namespace
    tbl_borders_xml = (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        "</w:tblBorders>"
    )

    tbl_borders = parse_xml(tbl_borders_xml)
    tbl_pr.append(tbl_borders)


def get_document_control_data(documentation):
    """Extract document control data for reporting"""
    return {
        "document_preparation": documentation.document_preparation,
        "document_title": documentation.document_title,
        "document_id": documentation.document_id,
        "document_version": documentation.document_version,
        "prepared_by": documentation.prepared_by,
        "reviewed_by": documentation.reviewed_by,
        "approved_by": documentation.approved_by,
        "released_by": documentation.released_by,
        "release_date": (
            documentation.release_date.strftime("%Y-%m-%d")
            if documentation.release_date
            else "N/A"
        ),
        "change_history": [
            {
                "version": change.version,
                "date": (
                    change.change_date.strftime("%Y-%m-%d")
                    if change.change_date
                    else "N/A"
                ),
                "remarks": change.remarks,
            }
            for change in documentation.change_history
        ],
        "distribution_list": [
            {
                "name": dist.name,
                "organization": dist.organization,
                "designation": dist.designation,
                "email": dist.email,
            }
            for dist in documentation.distribution_list
        ],
    }


def get_auditing_team_data(documentation):
    """Extract auditing team data for reporting"""
    return [
        {
            "name": member.name,
            "designation": member.designation,
            "email": member.email,
            "professional_qualifications": member.professional_qualifications,
            "listed_in_snapshot": member.listed_in_snapshot,
        }
        for member in documentation.audit_team
    ]


def get_tools_software_data(documentation):
    """Extract tools and software data for reporting"""
    return [
        {
            "tool_name": tool.tool_name,
            "version_control": tool.version_control,
            "license_type": tool.license_type,
        }
        for tool in documentation.tools_used
    ]


def add_documentation_sheets_to_excel(wb, documentation_data):
    """Add documentation sections as separate sheets in Excel"""
    print(f"DEBUG: Documentation data keys: {list(documentation_data.keys())}")
    
    # Document Control Sheet
    if documentation_data.get("document_control"):
        doc_data = documentation_data["document_control"]
        print(f"DEBUG: Document control data: {doc_data}")
        
        ws_doc = wb.create_sheet("Document Control")
        
        # Add basic document info
        ws_doc.append(["Document Control Information"])
        ws_doc.append([])  # Empty row for spacing
        
        # Document details in key-value format
        document_details = [
            ["Document Preparation", doc_data.get("document_preparation", "N/A")],
            ["Document Title", doc_data.get("document_title", "N/A")],
            ["Document ID", doc_data.get("document_id", "N/A")],
            ["Document Version", doc_data.get("document_version", "N/A")],
            ["Prepared By", doc_data.get("prepared_by", "N/A")],
            ["Reviewed By", doc_data.get("reviewed_by", "N/A")],
            ["Approved By", doc_data.get("approved_by", "N/A")],
            ["Released By", doc_data.get("released_by", "N/A")],
            ["Release Date", doc_data.get("release_date", "N/A")]
        ]
        
        for detail in document_details:
            ws_doc.append(detail)
        
        # Add change history if exists
        if doc_data.get("change_history"):
            ws_doc.append([])
            ws_doc.append(["Document Change History"])
            ws_doc.append(["Version", "Date", "Remarks"])
            for change in doc_data["change_history"]:
                ws_doc.append([change.get("version", "N/A"), change.get("date", "N/A"), change.get("remarks", "N/A")])
        
        # Add distribution list if exists
        if doc_data.get("distribution_list"):
            ws_doc.append([])
            ws_doc.append(["Document Distribution List"])
            ws_doc.append(["Name", "Organization", "Designation", "Email ID"])
            for dist in doc_data["distribution_list"]:
                ws_doc.append([
                    dist.get("name", "N/A"),
                    dist.get("organization", "N/A"),
                    dist.get("designation", "N/A"),
                    dist.get("email", "N/A")
                ])
        
        # Style the header
        for cell in ws_doc[1]:
            cell.font = Font(bold=True)
    
    # Introduction Sheet
    if documentation_data.get("introduction"):
        ws_intro = wb.create_sheet("Introduction")
        ws_intro.append(["Introduction"])
        ws_intro.append([])
        # Handle introduction text which might be long
        intro_text = documentation_data["introduction"]
        if isinstance(intro_text, str):
            # Split long text into multiple rows
            words = intro_text.split()
            lines = []
            current_line = ""
            for word in words:
                if len(current_line + " " + word) <= 100:  # Limit line length
                    current_line += " " + word if current_line else word
                else:
                    lines.append(current_line)
                    current_line = word
            if current_line:
                lines.append(current_line)
            
            for line in lines:
                ws_intro.append([line])
        else:
            ws_intro.append([str(intro_text)])
        
        # Style header
        ws_intro['A1'].font = Font(bold=True)
    
    # Engagement Scope Sheet
    if documentation_data.get("engagement_scope"):
        ws_scope = wb.create_sheet("Engagement Scope")
        ws_scope.append(["Engagement Scope"])
        ws_scope.append([])
        scope_text = documentation_data["engagement_scope"]
        if isinstance(scope_text, str):
            words = scope_text.split()
            lines = []
            current_line = ""
            for word in words:
                if len(current_line + " " + word) <= 100:
                    current_line += " " + word if current_line else word
                else:
                    lines.append(current_line)
                    current_line = word
            if current_line:
                lines.append(current_line)
            
            for line in lines:
                ws_scope.append([line])
        else:
            ws_scope.append([str(scope_text)])
        
        ws_scope['A1'].font = Font(bold=True)
    
    # Auditing Team Sheet
    if documentation_data.get("auditing_team"):
        ws_team = wb.create_sheet("Auditing Team")
        ws_team.append(["Auditing Team Members"])
        ws_team.append([])
        ws_team.append(["Name", "Designation", "Email", "Qualifications", "Listed in Snapshot"])
        
        for member in documentation_data["auditing_team"]:
            ws_team.append([
                member.get("name", "N/A"),
                member.get("designation", "N/A"),
                member.get("email", "N/A"),
                member.get("professional_qualifications", "N/A"),
                member.get("listed_in_snapshot", "N/A")
            ])
        
        # Style header row
        for cell in ws_team[3]:
            cell.font = Font(bold=True)
        ws_team['A1'].font = Font(bold=True)
    
    # Activities and Timelines Sheet
    if documentation_data.get("activities_timelines"):
        ws_activities = wb.create_sheet("Activities & Timelines")
        ws_activities.append(["Activities and Timelines"])
        ws_activities.append([])
        activities_text = documentation_data["activities_timelines"]
        if isinstance(activities_text, str):
            words = activities_text.split()
            lines = []
            current_line = ""
            for word in words:
                if len(current_line + " " + word) <= 100:
                    current_line += " " + word if current_line else word
                else:
                    lines.append(current_line)
                    current_line = word
            if current_line:
                lines.append(current_line)
            
            for line in lines:
                ws_activities.append([line])
        else:
            ws_activities.append([str(activities_text)])
        
        ws_activities['A1'].font = Font(bold=True)
    
    # Methodology and Criteria Sheet
    if documentation_data.get("methodology_criteria"):
        ws_methodology = wb.create_sheet("Methodology & Criteria")
        ws_methodology.append(["Methodology and Criteria"])
        ws_methodology.append([])
        methodology_text = documentation_data["methodology_criteria"]
        if isinstance(methodology_text, str):
            words = methodology_text.split()
            lines = []
            current_line = ""
            for word in words:
                if len(current_line + " " + word) <= 100:
                    current_line += " " + word if current_line else word
                else:
                    lines.append(current_line)
                    current_line = word
            if current_line:
                lines.append(current_line)
            
            for line in lines:
                ws_methodology.append([line])
        else:
            ws_methodology.append([str(methodology_text)])
        
        ws_methodology['A1'].font = Font(bold=True)
    
    # Tools and Software Sheet
    if documentation_data.get("tools_software"):
        ws_tools = wb.create_sheet("Tools & Software")
        ws_tools.append(["Tools and Software Used"])
        ws_tools.append([])
        ws_tools.append(["Tool Name", "Version", "License Type"])
        
        for tool in documentation_data["tools_software"]:
            ws_tools.append([
                tool.get("tool_name", "N/A"),
                tool.get("version_control", "N/A"),
                tool.get("license_type", "N/A")
            ])
        
        # Style header row
        for cell in ws_tools[3]:
            cell.font = Font(bold=True)
        ws_tools['A1'].font = Font(bold=True)
    
    print(f"DEBUG: Created {len(wb.sheetnames)} sheets in Excel workbook")

def process_audit_observations(
    doc_path, output_path, all_observations, documentation_data=None,clauses=None,clause_activities=None
):
    """
    Loads a DOCX document, adds documentation sections if selected,
    then adds a "Consolidated Audit Observations" heading with tables for each clause,
    and saves the document.
    """
    try:
        # Load the document
        document = Document(doc_path)
        print(f"Document '{doc_path}' loaded successfully.")
        print(f"Processing {len(all_observations)} consolidated clauses")

        # Add documentation sections if selected
        if documentation_data:
            print(f"Adding documentation sections: {list(documentation_data.keys())}")
            add_documentation_sections_to_word_simple(document, documentation_data)
        else:
            print("No documentation data to add")


          

        # Add the main "Consolidated Audit Observations" heading
        document.add_paragraph()  # Add some space before the main heading
        main_heading_paragraph = document.add_heading(
            "Consolidated Audit Observations", level=1
        )
        main_heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add introductory paragraph explaining the consolidated nature
        intro_paragraph = document.add_paragraph()
        intro_paragraph.add_run(
            "This report presents consolidated audit observations at the clause level, "
            "including comprehensive test procedures, observation summaries, findings, "
            "and recommendations aggregated from all applicable control activities."
        ).italic = True
        document.add_paragraph()  # Add spacing

        # Style the heading
        if main_heading_paragraph.runs:
            run = main_heading_paragraph.runs[0]
            run.font.size = Pt(18)
            try:
                run.font.color.rgb = RGBColor(0xB2, 0x22, 0x22)  # Firebrick color
            except Exception as e:
                print(f"Warning: Could not set exact heading color. {e}")

        # Add each consolidated clause observation as a separate table
        for i, obs in enumerate(all_observations):
            # Add clause heading
            clause_heading = document.add_heading(
                f"Clause {obs.get('Reference no.', 'N/A')}", level=2
            )

            # Use the enhanced consolidated table function
            add_consolidated_observation_table(document, obs)

            # Add the activities table for this clause if we have the data
            if clauses and clause_activities and i < len(clauses):
                clause = clauses[i]
                clause_id = clause.id
                activities_data = clause_activities.get(clause_id, [])
                if activities_data:
                    add_activities_table_to_word(document, activities_data, obs.get('Reference no.', 'N/A'))

                    # ADD PAGE BREAK AFTER ACTIVITIES TABLE (but not for last clause)
                    if i < len(all_observations) - 1:
                        document.add_page_break()
            
            else:
                # Add page break after every 2 clauses to avoid overly long pages
                if i < len(all_observations) - 1 and i % 2 == 1:
                    document.add_page_break()

        # Ensure the output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # Save the modified document
        document.save(output_path)
        print(f"Consolidated document successfully saved to '{output_path}'.")

    except Exception as e:
        print(f"An error occurred in process_audit_observations: {e}")
        import traceback

        traceback.print_exc()
        raise


def add_documentation_sections_to_word_simple(document, documentation_data):
    """Add documentation sections to Word document with borders and proper sizing"""

    # Document Control Section
    if documentation_data.get("document_control"):
        doc_data = documentation_data["document_control"]

        # Add section heading
        doc_heading = document.add_heading("Document Control", level=1)
        if doc_heading.runs:
            doc_heading.runs[0].font.color.rgb = RGBColor(0xB2, 0x22, 0x22)

        # Create a table for document control info
        table = document.add_table(rows=9, cols=2)
        table.autofit = False
        table.allow_autofit = False

        # Set optimized column widths
        table.columns[0].width = Inches(1.8)  # Field names
        table.columns[1].width = Inches(4.5)  # Values

        # Add borders to the table
        add_table_borders(table)

        # Populate document control table
        doc_fields = [
            ("Document Preparation", doc_data.get("document_preparation", "N/A")),
            ("Document Title", doc_data.get("document_title", "N/A")),
            ("Document ID", doc_data.get("document_id", "N/A")),
            ("Document Version", doc_data.get("document_version", "N/A")),
            ("Prepared By", doc_data.get("prepared_by", "N/A")),
            ("Reviewed By", doc_data.get("reviewed_by", "N/A")),
            ("Approved By", doc_data.get("approved_by", "N/A")),
            ("Released By", doc_data.get("released_by", "N/A")),
            ("Release Date", doc_data.get("release_date", "N/A")),
        ]

        for i, (field, value) in enumerate(doc_fields):
            row_cells = table.rows[i].cells

            # Truncate long values
            display_value = str(value)
            if len(display_value) > 80:
                display_value = display_value[:77] + "..."

            row_cells[0].text = field
            row_cells[1].text = display_value

            # Make field names bold with optimized font size
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        document.add_paragraph()  # Add space

        # Add Document Change History Table
        if doc_data.get("change_history"):
            add_change_history_table_with_borders(document, doc_data["change_history"])

        # Add Document Distribution List Table
        if doc_data.get("distribution_list"):
            add_distribution_list_table_with_borders(
                document, doc_data["distribution_list"]
            )
        
        # Add page break after Document Control section
        document.add_page_break()

    # Add other sections...
    if documentation_data.get("introduction"):
        add_rich_text_section(
            document, "Introduction", documentation_data["introduction"]
        )

    if documentation_data.get("engagement_scope"):
        add_rich_text_section(
            document, "Engagement Scope", documentation_data["engagement_scope"]
        )

    if documentation_data.get("activities_timelines"):
        add_rich_text_section(
            document,
            "Activities & Timelines",
            documentation_data["activities_timelines"],
        )

    if documentation_data.get("methodology_criteria"):
        add_rich_text_section(
            document,
            "Methodology & Criteria",
            documentation_data["methodology_criteria"],
        )

    if documentation_data.get("auditing_team"):
        add_auditing_team_table_with_borders(
            document, documentation_data["auditing_team"]
        )

    if documentation_data.get("tools_software"):
        add_tools_table_with_borders(document, documentation_data["tools_software"])

    if documentation_data.get('executive_summary'):
        add_executive_summary_to_word(document, documentation_data['executive_summary'])  
        document.add_page_break()  


def add_change_history_table_with_borders(document, change_history_data):
    """Add Document Change History table with borders and optimized widths"""
    # Add subheading for Change History
    change_heading = document.add_heading("Document Change History", level=2)
    if change_heading.runs:
        change_heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    if change_history_data:
        # Create table with headers
        table = document.add_table(rows=1, cols=3)
        table.autofit = False
        table.allow_autofit = False

        # Set optimized column widths
        table.columns[0].width = Inches(1.2)  # Version
        table.columns[1].width = Inches(1.2)  # Date
        table.columns[2].width = Inches(1.2)  # Remarks

        # Add borders to the table
        add_table_borders(table)

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Version"
        hdr_cells[1].text = "Date"
        hdr_cells[2].text = "Remarks"

        # Make header bold with smaller font
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Add data rows with text handling
        for change in change_history_data:
            row_cells = table.add_row().cells
            row_cells[0].text = change.get("version", "N/A")
            row_cells[1].text = change.get("date", "N/A")

            # Handle long remarks
            remarks = change.get("remarks", "N/A")
            if len(remarks) > 100:
                remarks = remarks[:97] + "..."
            row_cells[2].text = remarks

    document.add_paragraph()


def add_distribution_list_table_with_borders(document, distribution_data):
    """Add Document Distribution List table with borders and optimized widths"""
    # Add subheading for Distribution List
    dist_heading = document.add_heading("Document Distribution List", level=2)
    if dist_heading.runs:
        dist_heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    if distribution_data:
        # Create table with headers
        table = document.add_table(rows=1, cols=4)
        table.autofit = False
        table.allow_autofit = False

        # Set optimized column widths
        table.columns[0].width = Inches(1.3)  # Name
        table.columns[1].width = Inches(1.5)  # Organization
        table.columns[2].width = Inches(1.3)  # Designation
        table.columns[3].width = Inches(2.1)  # Email ID

        # Add borders to the table
        add_table_borders(table)

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Name"
        hdr_cells[1].text = "Organization"
        hdr_cells[2].text = "Designation"
        hdr_cells[3].text = "Email ID"

        # Make header bold with smaller font
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Add data rows with text optimization
        for dist in distribution_data:
            row_cells = table.add_row().cells

            # Name
            name = dist.get("name", "N/A")
            if len(name) > 20:
                name = name[:17] + "..."
            row_cells[0].text = name

            # Organization
            org = dist.get("organization", "N/A")
            if len(org) > 25:
                org = org[:22] + "..."
            row_cells[1].text = org

            # Designation
            designation = dist.get("designation", "N/A")
            if len(designation) > 20:
                designation = designation[:17] + "..."
            row_cells[2].text = designation

            # Email
            email = dist.get("email", "N/A")
            if len(email) > 30:
                email = email[:27] + "..."
            row_cells[3].text = email

    document.add_paragraph()


def add_auditing_team_table_with_borders(document, team_data):
    """Add auditing team table to Word document with borders and proper column widths"""
    heading = document.add_heading("Auditing Team", level=1)
    if heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(0xB2, 0x22, 0x22)

    if team_data:
        # Create table with headers
        table = document.add_table(rows=1, cols=5)
        table.autofit = False
        table.allow_autofit = False

        # Set optimized column widths to fit within page margins
        # Total page width is about 6.5 inches, so distribute accordingly
        table.columns[0].width = Inches(1.2)  # Name (reduced)
        table.columns[1].width = Inches(1.2)  # Designation (reduced)
        table.columns[2].width = Inches(1.5)  # Email (reduced)
        table.columns[3].width = Inches(1.8)  # Qualifications (reduced)
        table.columns[4].width = Inches(1.0)  # Listed in Snapshot (reduced)

        # Add borders to the table
        add_table_borders(table)

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Name"
        hdr_cells[1].text = "Designation"
        hdr_cells[2].text = "Email"
        hdr_cells[3].text = "Qualifications"
        hdr_cells[4].text = "Listed in Snapshot"

        # Make header bold and adjust font size
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)  # Smaller font for headers

        # Add data rows with optimized text wrapping
        for member in team_data:
            row_cells = table.add_row().cells

            # Name - truncate if too long
            name = member.get("name", "N/A")
            if len(name) > 25:
                name = name[:22] + "..."
            row_cells[0].text = name

            # Designation - truncate if too long
            designation = member.get("designation", "N/A")
            if len(designation) > 25:
                designation = designation[:22] + "..."
            row_cells[1].text = designation

            # Email - truncate if too long
            email = member.get("email", "N/A")
            if len(email) > 30:
                email = email[:27] + "..."
            row_cells[2].text = email

            # Qualifications - handle long text with wrapping
            qualifications = member.get("professional_qualifications", "N/A")
            if len(qualifications) > 40:
                # Split into multiple lines if too long
                words = qualifications.split()
                lines = []
                current_line = ""
                for word in words:
                    if len(current_line + " " + word) <= 35:
                        current_line += " " + word
                    else:
                        if current_line:
                            lines.append(current_line.strip())
                        current_line = word
                if current_line:
                    lines.append(current_line.strip())
                qualifications = "\n".join(lines)
            row_cells[3].text = qualifications

            # Listed in Snapshot - simple yes/no
            row_cells[4].text = member.get("listed_in_snapshot", "N/A")

            # Set smaller font size for data cells to fit better
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Smaller font for data

    document.add_paragraph()


def add_tools_table_with_borders(document, tools_data):
    """Add tools and software table to Word document with borders and optimized widths"""
    heading = document.add_heading("Tools & Software", level=1)
    if heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(0xB2, 0x22, 0x22)

    if tools_data:
        table = document.add_table(rows=1, cols=3)
        table.autofit = False
        table.allow_autofit = False

        # Set optimized column widths
        table.columns[0].width = Inches(2.2)  # Tool Name
        table.columns[1].width = Inches(1.3)  # Version
        table.columns[2].width = Inches(1.2)  # License Type

        # Add borders to the table
        add_table_borders(table)

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Tool Name"
        hdr_cells[1].text = "Version"
        hdr_cells[2].text = "License Type"

        # Make header bold with smaller font
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Add data rows with text optimization
        for tool in tools_data:
            row_cells = table.add_row().cells

            # Tool Name
            tool_name = tool.get("tool_name", "N/A")
            if len(tool_name) > 35:
                tool_name = tool_name[:32] + "..."
            row_cells[0].text = tool_name

            # Version
            version = tool.get("version_control", "N/A")
            if len(version) > 20:
                version = version[:17] + "..."
            row_cells[1].text = version

            # License Type
            row_cells[2].text = tool.get("license_type", "N/A")

    document.add_paragraph()
    document.add_page_break() 


def add_rich_text_section(document, title, content):
    """Add a rich text section to Word document"""
    heading = document.add_heading(title, level=1)
    if heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(0xB2, 0x22, 0x22)

    # Add the content (you might want to parse HTML here)
    paragraph = document.add_paragraph()
    # Simple text addition - you can enhance this to handle HTML formatting
    clean_content = clean_html_preserve_format(content) if content else "N/A"
    paragraph.add_run(clean_content)

    document.add_paragraph()  # Add space
    document.add_page_break()  # Add page break after section


def add_auditing_team_table(document, team_data):
    """Add auditing team table to Word document"""
    heading = document.add_heading("Auditing Team", level=1)
    if heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(0xB2, 0x22, 0x22)

    if team_data:
        table = document.add_table(rows=1, cols=5)
        table.autofit = False
        table.allow_autofit = False

        # Set column widths
        table.columns[0].width = Inches(1.5)  # Name
        table.columns[1].width = Inches(1.5)  # Designation
        table.columns[2].width = Inches(2.0)  # Email
        table.columns[3].width = Inches(2.0)  # Qualifications
        table.columns[4].width = Inches(1.5)  # Listed in Snapshot

        # Add borders manually
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:tcBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'
                    )
                )

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Name"
        hdr_cells[1].text = "Designation"
        hdr_cells[2].text = "Email"
        hdr_cells[3].text = "Qualifications"
        hdr_cells[4].text = "Listed in Snapshot"

        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for member in team_data:
            row_cells = table.add_row().cells
            row_cells[0].text = member.get("name", "N/A")
            row_cells[1].text = member.get("designation", "N/A")
            row_cells[2].text = member.get("email", "N/A")
            row_cells[3].text = member.get("professional_qualifications", "N/A")
            row_cells[4].text = member.get("listed_in_snapshot", "N/A")

            # Add borders to new row
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:tcBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'
                    )
                )

    document.add_paragraph()  # Add space
    document.add_page_break() 


def add_tools_table(document, tools_data):
    """Add tools and software table to Word document"""
    heading = document.add_heading("Tools & Software", level=1)
    if heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(0xB2, 0x22, 0x22)

    if tools_data:
        table = document.add_table(rows=1, cols=3)
        table.autofit = False
        table.allow_autofit = False

        # Set column widths
        table.columns[0].width = Inches(2.5)  # Tool Name
        table.columns[1].width = Inches(1.5)  # Version
        table.columns[2].width = Inches(1.5)  # License Type

        # Add borders manually
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:tcBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'
                    )
                )

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Tool Name"
        hdr_cells[1].text = "Version"
        hdr_cells[2].text = "License Type"

        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for tool in tools_data:
            row_cells = table.add_row().cells
            row_cells[0].text = tool.get("tool_name", "N/A")
            row_cells[1].text = tool.get("version_control", "N/A")
            row_cells[2].text = tool.get("license_type", "N/A")

            # Add borders to new row
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:tcBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'
                    )
                )

    document.add_paragraph()  # Add space


def add_change_history_table(document, change_history_data):
    """Add Document Change History table to Word document"""
    # Add subheading for Change History
    change_heading = document.add_heading("Document Change History", level=2)
    if change_heading.runs:
        change_heading.runs[0].font.color.rgb = RGBColor(
            0x1F, 0x4E, 0x79
        )  # Dark blue color

    if change_history_data:
        # Create table with headers
        table = document.add_table(rows=1, cols=3)
        table.autofit = False
        table.allow_autofit = False

        # Set column widths
        table.columns[0].width = Inches(1.5)  # Version
        table.columns[1].width = Inches(1.5)  # Date
        table.columns[2].width = Inches(3.0)  # Remarks

        # Add borders manually
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:tcBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'
                    )
                )

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Version"
        hdr_cells[1].text = "Date"
        hdr_cells[2].text = "Remarks"

        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)

        # Add data rows
        for change in change_history_data:
            row_cells = table.add_row().cells
            row_cells[0].text = change.get("version", "N/A")
            row_cells[1].text = change.get("date", "N/A")
            row_cells[2].text = change.get("remarks", "N/A")

            # Add borders to new row
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:tcBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'
                    )
                )

    document.add_paragraph()  # Add space after the table


def add_distribution_list_table(document, distribution_data):
    """Add Document Distribution List table to Word document"""
    # Add subheading for Distribution List
    dist_heading = document.add_heading("Document Distribution List", level=2)
    if dist_heading.runs:
        dist_heading.runs[0].font.color.rgb = RGBColor(
            0x1F, 0x4E, 0x79
        )  # Dark blue color

    if distribution_data:
        # Create table with headers
        table = document.add_table(rows=1, cols=4)
        table.autofit = False
        table.allow_autofit = False

        # Set column widths
        table.columns[0].width = Inches(1.5)  # Name
        table.columns[1].width = Inches(1.5)  # Organization
        table.columns[2].width = Inches(1.5)  # Designation
        table.columns[3].width = Inches(2.5)  # Email ID

        # Add borders manually
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:tcBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'
                    )
                )

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Name"
        hdr_cells[1].text = "Organization"
        hdr_cells[2].text = "Designation"
        hdr_cells[3].text = "Email ID"

        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)

        # Add data rows
        for dist in distribution_data:
            row_cells = table.add_row().cells
            row_cells[0].text = dist.get("name", "N/A")
            row_cells[1].text = dist.get("organization", "N/A")
            row_cells[2].text = dist.get("designation", "N/A")
            row_cells[3].text = dist.get("email", "N/A")

            # Add borders to new row
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:tcBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'
                    )
                )

    document.add_paragraph()  # Add space after the table


def extract_content(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return extract_from_pdf(path)
        elif ext == ".docx":
            return extract_from_docx(path)

        elif ext in [".xlsx", ".xls"]:
            return extract_from_excel(path)
        else:
            return "Unsupported file format."
    except Exception as e:
        return f"Error during extraction: {str(e)}"


def add_table_borders_simple(table):
    """Simple method to add borders to table using python-docx built-in methods"""
    # This method uses the existing table style but ensures borders are visible
    for row in table.rows:
        for cell in row.cells:
            # Set cell margins to make content look better
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)


def extract_from_excel(path):
    try:
        ext = os.path.splitext(path)[1].lower()
        text_content = []

        if ext == ".xlsx":
            # For xlsx files
            from openpyxl import load_workbook

            wb = load_workbook(filename=path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                text_content.append(f"--- Sheet: {sheet_name} ---")
                ws = wb[sheet_name]

                for row in ws.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    text_content.append("\t".join(row_text))

        elif ext == ".xls":
            # For xls files
            import xlrd

            wb = xlrd.open_workbook(path)
            for sheet_name in wb.sheet_names():
                text_content.append(f"--- Sheet: {sheet_name} ---")
                sheet = wb.sheet_by_name(sheet_name)

                for row_idx in range(sheet.nrows):
                    row = sheet.row(row_idx)
                    row_text = [str(cell.value) for cell in row]
                    text_content.append("\t".join(row_text))

        return "\n".join(text_content) or "[No extractable data found in Excel file]"

    except Exception as e:
        return f"Error extracting Excel content: {str(e)}"


def extract_from_pdf(path):
    text = ""
    doc = fitz.open(path)
    for page in doc:
        text += page.get_text()
    return text.strip() or "[No extractable text found in PDF]"


def extract_from_docx(path):
    doc = Document(path)
    return (
        "\n".join([para.text for para in doc.paragraphs])
        or "[No extractable text in DOCX]"
    )


def generate_chat_output(prompt):
    try:
        # Define the comprehensive system prompt
        system_prompt = """You are a compliance reviewer generating structured evidence assessments for enterprise audit software.

ROLE: Internal audit evidence assessor. NOT a legal advisor, consultant, or narrative writer.

OUTPUT FORMAT — ALWAYS follow this exact structure. Return markdown only:

### Document Reviewed
[Document name and entity name]

### Objective
[One sentence: what this control/activity is verifying]

### Evidence Coverage

| Requirement | Evidence Found | Policy / Document Reference | Coverage Status |
|---|---|---|---|
| [Control area] | [Specific finding — one line] | [Section/Page] | Covered / Partial / Not Evident |

(Minimum 4 rows, maximum 10 rows. Only control areas relevant to the activity objective.)

### Conclusion
[2-4 sentences maximum. Factual. What was assessed and whether objective is addressed.]

FORBIDDEN PHRASES — Never use:
overall, collectively, robustly, robust, comprehensively, comprehensive, holistic,
effectively aligns, effectively, substantively, multiple sections, various sections,
strong framework, strong, concerted effort, adequately demonstrates, adequate, seamlessly,
appropriately reflects, demonstrates commitment, rigorous, rigorously, thorough, thoroughly,
ensures, ensuring, highlight, highlights, underscores, emphasizes, it is evident,
demonstrates, showcase, paramount, pivotal, crucial, well-established, well-defined,
clearly demonstrates, strongly, robust framework, sound framework, solid framework

STRICT RULES:
- Maximum 700 words total
- No long paragraphs anywhere
- No narrative essays
- No repetition of same point
- No findings, recommendations, observations, ratings, or gap assessments
- No hallucinated sections or invented references
- Markdown table format only for Evidence Coverage
- Return markdown only — no preamble, no explanations, no chain-of-thought

PREFERRED WORDING:
Use: Defines, Documents, Specifies, Includes, Describes, Outlines, States, Lists
Avoid: adjective-heavy, AI-sounding language

TRACEABILITY:
- Always reference section/page/clause numbers where available
- If unclear: write "Not clearly evidenced" and mark as "Not Evident"
- Never hallucinate coverage

ENTITY MATCHING:
- Check document entity name matches auditee
- JK Bank, JK Bank Ltd = same entity
- If entity name present in document = ADMISSIBLE
- Do NOT mark as inadmissible if entity name is in document title or content

DO NOT summarize entire document. Only map evidence relevant to activity objective."""
        
        # Log the system prompt being sent
        current_app.logger.info("=" * 80)
        current_app.logger.info("SYSTEM PROMPT BEING SENT TO AI:")
        current_app.logger.info(system_prompt)
        current_app.logger.info("=" * 80)
        current_app.logger.info("USER PROMPT (first 1000 chars):")
        current_app.logger.info(prompt[:1000])
        current_app.logger.info("=" * 80)
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": system_prompt,
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.15,
            top_p=0.2,
            frequency_penalty=0.6,
            presence_penalty=0.0,
            max_tokens=1200,
        )
        
        result = response.choices[0].message.content.strip()
        
        # Log the AI response
        current_app.logger.info("=" * 80)
        current_app.logger.info("AI RESPONSE (first 2000 chars):")
        current_app.logger.info(result[:2000])
        current_app.logger.info("=" * 80)
        
        # More robust JSON extraction
        if "```json" in result:
            start_idx = result.find("```json") + 7
            end_idx = result.find("```", start_idx)
            if end_idx != -1:
                result = result[start_idx:end_idx].strip()
        elif "```" in result:
            start_idx = result.find("```") + 3
            end_idx = result.find("```", start_idx)
            if end_idx != -1:
                result = result[start_idx:end_idx].strip()
        
        try:
            json.loads(result)
            return result
        except json.JSONDecodeError:
            start_idx = result.find('{')
            end_idx = result.rfind('}') + 1
            if start_idx != -1 and end_idx != -1:
                json_str = result[start_idx:end_idx]
                json.loads(json_str)
                return json_str
            
            current_app.logger.error(f"AI returned invalid JSON: {result[:500]}")
            return '{"grouped_evidences": [], "error": "Invalid JSON response from AI"}'
            
    except Exception as e:
        current_app.logger.error(f"Error in generate_chat_output: {str(e)}")
        return '{"grouped_evidences": [], "error": "AI service error"}'



def get_executive_summary_data(documentation, project, clause_statistics, severity_stats, evidence_stats):
    """
    Extract executive summary data from documentation and project for the report
    Matches the exact structure shown in documentation.html
    """
    if not documentation:
        return None
    
    # Get the narrative executive summary (if edited by user)
    executive_summary_narrative = documentation.executive_summary
    
    # Get locations data exactly as stored in documentation
    locations_data = []
    if documentation and documentation.locations_data:
        locations_data = documentation.locations_data
    elif project.client_rel:
        # Fallback to building from project if no documentation locations
        if hasattr(project.client_rel, 'addresses') and project.client_rel.addresses:
            for addr in project.client_rel.addresses:
                # Build complete address string
                address_parts = []
                if hasattr(addr, 'address_line1') and addr.address_line1:
                    address_parts.append(addr.address_line1)
                if hasattr(addr, 'address_line2') and addr.address_line2:
                    address_parts.append(addr.address_line2)
                if hasattr(addr, 'city') and addr.city:
                    address_parts.append(addr.city)
                if hasattr(addr, 'state') and addr.state:
                    address_parts.append(addr.state)
                if hasattr(addr, 'country') and addr.country:
                    address_parts.append(addr.country)
                if hasattr(addr, 'postal_code') and addr.postal_code:
                    address_parts.append(addr.postal_code)
                
                full_address = ", ".join(filter(None, address_parts))
                
                locations_data.append({
                    'city': addr.city if hasattr(addr, 'city') else '',
                    'country': addr.country if hasattr(addr, 'country') else '',
                    'address': full_address
                })
    
    # Get departments data exactly as stored in documentation
    departments_data = []
    if documentation and documentation.departments_data:
        departments_data = documentation.departments_data
    elif project.departments:
        seen = set()
        for dept in project.departments:
            if dept.department_name not in seen:
                seen.add(dept.department_name)
                departments_data.append({
                    'name': dept.department_name
                })
    elif project.primary_department:
        departments_data.append({
            'name': project.primary_department.department_name
        })
    
    # Get audit organization name
    audit_org_name = None
    if documentation.auditor_profile_id:
        audit_org = AuditOrganization.query.get(documentation.auditor_profile_id)
        if audit_org:
            audit_org_name = audit_org.firm_name
    
    # Get guideline name
    guideline_name = get_guideline_names(project)
    guideline_release_date = get_guideline_release_date(project)
    
    return {
        'narrative': executive_summary_narrative,
        'locations': locations_data,
        'departments': departments_data,
        'audit_org_name': audit_org_name or 'Audit Organization Name',
        'client_name': project.client_rel.name if project.client_rel else 'Client Name',
        'assessment_start': project.assesment_start_date.strftime('%d %B %Y') if project.assesment_start_date else 'Assessment Period',
        'assessment_end': project.assesment_end_date.strftime('%d %B %Y') if project.assesment_end_date else 'Assessment Period',
        'guideline_name': guideline_name or 'Guideline Name',
        'guideline_release_date': guideline_release_date or 'Guideline release date',
        'clause_statistics': clause_statistics,
        'severity_stats': severity_stats,
        'evidence_stats': evidence_stats
    }



def get_guideline_names(project):
    """Helper to get guideline names as string - matches documentation.html"""
    guidelines_list = []
    if project.guidelines_rel:
        guideline = project.guidelines_rel
        if guideline.guideline_data:
            doc_details = guideline.guideline_data.get('DocumentDetails', {})
            guidelines_list.append(doc_details.get('DocumentName', 'Applicable Guidelines'))
    elif project.project_guidelines:
        for pg in project.project_guidelines:
            if pg.guideline_data:
                doc_details = pg.guideline_data.get('DocumentDetails', {})
                guidelines_list.append(doc_details.get('DocumentName', 'Applicable Guidelines'))
    
    return ", ".join(guidelines_list) if guidelines_list else 'Applicable Guidelines'

def get_guideline_release_date(project):
    """Helper to get guideline release date"""
    guidelines_list = []
    if project.guidelines_rel:
        guideline = project.guidelines_rel
        if guideline.guideline_data:
            doc_details = guideline.guideline_data.get('DocumentDetails', {})
            guidelines_list.append(doc_details.get('IssuanceDate') or 'Not Specified')
    elif project.project_guidelines:
        for pg in project.project_guidelines:
            if pg.guideline_data:
                doc_details = pg.guideline_data.get('DocumentDetails', {})
                guidelines_list.append(doc_details.get('IssuanceDate') or 'Not Specified')
    
    # Filter out 'Not Specified' and join
    valid_dates = [d for d in guidelines_list if d != 'Not Specified']
    return ", ".join(valid_dates) if valid_dates else 'Not Specified'


def get_audit_org_name(auditor_profile_id):
    """Helper to get audit organization name"""
    if auditor_profile_id:
        audit_org = AuditOrganization.query.get(auditor_profile_id)
        if audit_org:
            return audit_org.firm_name
    return 'Audit Organization Name'

def add_executive_summary_to_word(document, executive_summary_data):
    """
    Add executive summary section to Word document with exact table structure from documentation.html
    Uses manual border formatting instead of 'Table Grid' style
    """
    if not executive_summary_data:
        return
    
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    
    def set_cell_borders(cell):
        """Add borders to a table cell"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Create border elements
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Black color
            
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            tcBorders.append(border)
    
    # Add Executive Summary heading
    heading = document.add_heading('Executive Summary', level=1)

    # Set heading color to red
    for run in heading.runs:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
    
    # Add the opening paragraph exactly as in documentation.html
    p = document.add_paragraph()
    audit_org = executive_summary_data.get('audit_org_name', 'Audit Organization Name')
    client = executive_summary_data.get('client_name', 'Client Name')
    start_date = executive_summary_data.get('assessment_start', 'Assessment Period')
    end_date = executive_summary_data.get('assessment_end', 'Assessment Period')
    guideline = executive_summary_data.get('guideline_name', 'Guideline Name')
    release_date = executive_summary_data.get('guideline_release_date', 'Guideline release date')
    
    p.add_run(f"{audit_org} ").bold = True
    p.add_run("conducted an audit of ")
    p.add_run(f"{client} ").bold = True
    p.add_run("from ")
    p.add_run(f"{start_date} ").bold = True
    p.add_run("to ")
    p.add_run(f"{end_date}").bold = True
    p.add_run(", as per the ")
    p.add_run(f"{guideline}").bold = True
    p.add_run(" dated ")
    p.add_run(f"{release_date}").bold = True
    p.add_run(".")
    
    document.add_paragraph()
    
    # Add locations covered text
    p = document.add_paragraph()
    p.add_run("The below locations and departments were covered under the scope of the audit:")
    
    document.add_paragraph()
    
    # Add Locations table - EXACT structure from documentation.html
    locations = executive_summary_data.get('locations', [])
    if locations:
        p = document.add_paragraph()
        p.add_run("List of Locations:").bold = True
        
        # Create table with 3 columns: Sr. No, Location, Address
        table = document.add_table(rows=len(locations) + 1, cols=3)
        table.autofit = False
        
        # Set column widths
        table.columns[0].width = Inches(0.8)  # Sr. No
        table.columns[1].width = Inches(2.0)  # Location
        table.columns[2].width = Inches(4.0)  # Address
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Sr. No"
        header_cells[1].text = "Location"
        header_cells[2].text = "Address"
        
        # Style header row - bold and add borders
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            set_cell_borders(cell)
        
        # Add location data rows
        for i, location in enumerate(locations, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(i)
            row_cells[1].text = f"{location.get('city', '')}, {location.get('country', '')}"
            row_cells[2].text = location.get('address', '')
            
            # Add borders to data cells
            for cell in row_cells:
                set_cell_borders(cell)
        
        document.add_paragraph()
    
    # Add Departments table - EXACT structure from documentation.html
    departments = executive_summary_data.get('departments', [])
    if departments:
        p = document.add_paragraph()
        p.add_run("List of Departments:").bold = True
        
        # Create table with 2 columns: Sr. No, Department Name
        table = document.add_table(rows=len(departments) + 1, cols=2)
        table.autofit = False
        
        # Set column widths
        table.columns[0].width = Inches(0.8)  # Sr. No
        table.columns[1].width = Inches(5.0)  # Department Name
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Sr. No"
        header_cells[1].text = "Department Name"
        
        # Style header row - bold and add borders
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            set_cell_borders(cell)
        
        # Add department data rows
        for i, dept in enumerate(departments, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(i)
            row_cells[1].text = dept.get('name', '')
            
            # Add borders to data cells
            for cell in row_cells:
                set_cell_borders(cell)
        
        document.add_paragraph()
    
    # Add the audit summary paragraph from the bottom of the dashboard
    clause_stats = executive_summary_data.get('clause_statistics', {})
    evidence_stats = executive_summary_data.get('evidence_stats', {})
    
    if clause_stats and evidence_stats:
        total_clauses = clause_stats.get('total_clauses', 0)
        applicable = clause_stats.get('applicability', {}).get('applicable', 0)
        with_evidence = evidence_stats.get('with_evidence', 0)
        compliant = clause_stats.get('compliance', {}).get('compliant', 0)
        partially_compliant = clause_stats.get('compliance', {}).get('partially_compliant', 0)
        non_compliant = clause_stats.get('compliance', {}).get('non_compliant', 0)
        
        p = document.add_paragraph()
        p.add_run("Out of the full list of ").bold = False
        p.add_run(str(total_clauses)).bold = True
        p.add_run(" clauses in the guideline, ")
        p.add_run(str(applicable)).bold = True
        p.add_run(" were scoped-in. Evidence and information was provided for assessment of ")
        p.add_run(str(with_evidence)).bold = True
        p.add_run(" clauses, out of which ")
        p.add_run(str(compliant)).bold = True
        p.add_run(" were found to be fully compliant.")
        
        p = document.add_paragraph()
        p.add_run("The audit identified ").bold = False
        p.add_run(str(partially_compliant + non_compliant)).bold = True
        p.add_run(" clauses to be partially compliant or non-compliant.")
    
    # Add spacing
    document.add_paragraph()

def get_clause_applicability_status(clause_id):
    """Get the applicability status of a clause"""
    try:
        clause = ProjectClause.query.get(clause_id)
        if clause:
            return "Applicable" if clause.applicability else "Not Applicable"
        return "N/A"
    except Exception as e:
        current_app.logger.error(f"Error getting applicability for clause {clause_id}: {e}")
        return "N/A"

def get_clause_assessment_status(clause_id):
    """Get the assessment status of a clause"""
    try:
        clause = ProjectClause.query.get(clause_id)
        if clause:
            return clause.assessment_status if clause.assessment_status else "To Be Assessed"
        return "To Be Assessed"
    except Exception as e:
        current_app.logger.error(f"Error getting assessment status for clause {clause_id}: {e}")
        return "To Be Assessed"


def get_clause_evidence_availability(clause_id):
    """
    Determine if all applicable activities for a clause have evidence received.
    Returns "Yes" if 100% of applicable activities have evidence, "No" otherwise.
    """
    try:
        # Get all applicable control activities for this clause
        control_activities = (
            db.session.query(ProjectControlActivity)
            .join(
                ProjectComplianceActivity,
                ProjectControlActivity.project_compliance_activity_id == ProjectComplianceActivity.id
            )
            .filter(ProjectComplianceActivity.project_clause_id == clause_id)
            .filter(ProjectComplianceActivity.applicability == True)  # Only applicable activities
            .all()
        )
        
        # If there are no applicable activities, return "N/A" or "No"
        if not control_activities:
            return "N/A"
        
        # Count activities with evidence received
        activities_with_evidence = 0
        total_applicable = len(control_activities)
        
        for activity in control_activities:
            # Check if evidence is received (admissible and strong)
            evidence_received = (
                activity.evidence_admissibility_decision == "Yes" and 
                activity.evidence_quality_rating == "STRONG"
            )
            if evidence_received:
                activities_with_evidence += 1
        
        # Calculate percentage
        evidence_percentage = (activities_with_evidence / total_applicable) * 100
        
        # Return "Yes" only if 100% of activities have evidence
        return "Yes" if evidence_percentage == 100 else "No"
        
    except Exception as e:
        current_app.logger.error(f"Error getting evidence availability for clause {clause_id}: {e}")
        return "Error"


def get_activities_for_clause_table(clause_id):
    """
    Fetch all applicable activities for a clause with their evidence admissibility,
    evidence quality, and required effectiveness data for report tables.
    
    Args:
        clause_id: The ID of the clause
    
    Returns:
        List of dictionaries containing activity data for the table
    """
    try:
        
        
        # Get all applicable control activities for this clause
        control_activities = (
            db.session.query(ProjectControlActivity)
            .join(
                ProjectComplianceActivity,
                ProjectControlActivity.project_compliance_activity_id == ProjectComplianceActivity.id
            )
            .filter(ProjectComplianceActivity.project_clause_id == clause_id)
            .filter(ProjectComplianceActivity.applicability == True)  # Only applicable activities
            .all()
        )
        
        activities_data = []
        
        for activity in control_activities:
            # Determine evidence received status
            evidence_received = (
                activity.evidence_admissibility_decision == "Yes" and 
                activity.evidence_quality_rating == "STRONG"
            )
            
            # Format evidence quality with appropriate display
            evidence_quality_display = activity.evidence_quality_rating or "Not Rated"
            if evidence_quality_display == "STRONG":
                evidence_quality_display = "STRONG ✓"
            elif evidence_quality_display == "ADEQUATE":
                evidence_quality_display = "ADEQUATE"
            elif evidence_quality_display == "WEAK":
                evidence_quality_display = "WEAK ⚠"
            elif evidence_quality_display == "INADMISSIBLE":
                evidence_quality_display = "INADMISSIBLE ✗"
            
            # Format admissibility decision
            admissibility_display = activity.evidence_admissibility_decision or "Not Evaluated"
            if admissibility_display == "Yes":
                admissibility_display = "Admissible ✓"
            elif admissibility_display == "No":
                admissibility_display = "Inadmissible ✗"
            
            # Format required effectiveness dimensions
            def format_effectiveness(value):
                if value == "yes":
                    return "Required ✓"
                elif value == "no":
                    return "Not Required"
                else:
                    return "N/A"
            
            # Get compliance status with proper formatting
            compliance_status = activity.compliant_status or "To be Assessed"
            
            # Get severity with appropriate label
            severity = activity.overall_severity_classification or "Not Classified"
            
            # Get parent compliance activity description
            parent_activity = activity.project_compliance_activity
            activity_name = parent_activity.activity_description if parent_activity else activity.activity_description
            
            # Count evidence files (if any)
            evidence_count = len(activity.submitted_evidences) if activity.submitted_evidences else 0
            
            # Get evidence file details (for tooltips)
            evidence_files = []
            if activity.submitted_evidences:
                for evidence in activity.submitted_evidences:
                    evidence_files.append({
                        'item': evidence.item,
                        'category': evidence.category,
                        'has_file': bool(evidence.evidence_file_path)
                    })
            
            activity_entry = {
                'activity_code': activity.activity_code,
                'activity_name': activity_name,
                'activity_description': activity.activity_description,
                'compliance_status': compliance_status,
                'severity': severity,
                'evidence_admissibility': admissibility_display,
                'evidence_quality': evidence_quality_display,
                'evidence_received': evidence_received,
                'evidence_count': evidence_count,
                'evidence_files': evidence_files,
                'required_effectiveness_design': format_effectiveness(activity.required_effectiveness_design),
                'required_effectiveness_implementation': format_effectiveness(activity.required_effectiveness_implementation),
                'required_effectiveness_operating': format_effectiveness(activity.required_effectiveness_operating),
                'owner': activity.owner,
                'control_type': activity.control_type,
                'frequency': activity.frequency,
                'has_findings': compliance_status in ['Non-Compliant', 'Partially Compliant'],
                'findings': activity.findings if activity.findings else None,
                'recommendations': activity.recommendations if activity.recommendations else None,
                'auditor_observation': activity.auditor_observation,
                'detailed_testing_results': activity.detailed_control_testing_results
            }
            
            activities_data.append(activity_entry)
        
        # Sort activities by activity_code (natural sorting)
        def natural_sort_key(a):
            text = a.get('activity_code', '')
            if not text:
                return [float('inf')]
            return [
                int(part) if part.isdigit() else part.lower()
                for part in re.split(r'(\d+)', str(text))
            ]
        
        activities_data = sorted(activities_data, key=natural_sort_key)
        
        current_app.logger.info(f"Found {len(activities_data)} applicable activities for clause {clause_id}")
        return activities_data
        
    except Exception as e:
        current_app.logger.error(f"Error getting activities for clause table: {e}")
        return []


def add_activities_table_to_word(document, activities_data, clause_ref):
    """
    Add activities table to Word document for a specific clause
    Optimized to fit within page margins with better column sizing
    """
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    if not activities_data:
        return
    
    def set_cell_borders(cell):
        """Add borders to a table cell"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    
    def set_table_borders(table):
        """Add borders to the entire table"""
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        
        tbl_borders_xml = (
            '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>'
        )
        
        tbl_borders = parse_xml(tbl_borders_xml)
        tbl_pr.append(tbl_borders)
    

    # ADD PAGE BREAK BEFORE ACTIVITIES TABLE
    document.add_page_break()
    
    # Add subheading for activities
    heading = document.add_heading(f'Applicable Activities for Clause {clause_ref}', level=3)
    if heading.runs:
        heading.runs[0].font.size = Pt(12)
        heading.runs[0].font.bold = True
    
    # OPTION 1: Split into two tables (Recommended - Best readability)
    # First table: Core information
    headers1 = [
        "Activity Code",
        "Activity Name",
        "Status",
        "Severity"
    ]
    
    table1 = document.add_table(rows=1, cols=len(headers1))
    table1.autofit = False
    table1.allow_autofit = False
    
    # Set column widths for first table (total ~6 inches)
    column_widths1 = [0.8, 2.5, 1.2, 1.2]
    for i, width in enumerate(column_widths1):
        table1.columns[i].width = Inches(width)
    
    set_table_borders(table1)
    
    # Add headers for first table
    header_row1 = table1.rows[0]
    for i, header in enumerate(headers1):
        cell = header_row1.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(cell)
    
    # Add data rows for first table
    for activity in activities_data:
        row = table1.add_row()
        
        # Activity Code
        row.cells[0].text = activity['activity_code']
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(row.cells[0])
        
        # Activity Name (truncated more aggressively)
        activity_name = activity['activity_name']
        if len(activity_name) > 60:
            activity_name = activity_name[:57] + "..."
        row.cells[1].text = activity_name
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(row.cells[1])
        
        # Compliance Status with color
        status_cell = row.cells[2]
        status_cell.text = activity['compliance_status']
        status_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in status_cell.paragraphs:
            if activity['compliance_status'] == 'Compliant':
                paragraph.runs[0].font.color.rgb = RGBColor(0, 100, 0)
            elif activity['compliance_status'] == 'Partially Compliant':
                paragraph.runs[0].font.color.rgb = RGBColor(156, 87, 0)
            elif activity['compliance_status'] == 'Non-Compliant':
                paragraph.runs[0].font.color.rgb = RGBColor(156, 0, 6)
        set_cell_borders(status_cell)
        
        # Severity with color
        severity_cell = row.cells[3]
        severity_cell.text = activity['severity']
        severity_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in severity_cell.paragraphs:
            if activity['severity'] == 'Critical':
                paragraph.runs[0].font.color.rgb = RGBColor(156, 0, 6)
                paragraph.runs[0].font.bold = True
            elif activity['severity'] == 'Major':
                paragraph.runs[0].font.color.rgb = RGBColor(156, 87, 0)
            elif activity['severity'] == 'No findings noted':
                paragraph.runs[0].font.color.rgb = RGBColor(0, 100, 0)
        set_cell_borders(severity_cell)
    
    document.add_paragraph()  # Add spacing
    
    # Second table: Evidence and Effectiveness information
    headers2 = [
        "Evidence\nAdmissibility",
        "Evidence\nQuality",
        "Design\nEffectiveness",
        "Implementation\nEffectiveness",
        "Operating\nEffectiveness"
    ]
    
    table2 = document.add_table(rows=1, cols=len(headers2))
    table2.autofit = False
    table2.allow_autofit = False
    
    # Set column widths for second table (total ~6 inches)
    column_widths2 = [1.2, 1.2, 1.2, 1.2, 1.2]
    for i, width in enumerate(column_widths2):
        table2.columns[i].width = Inches(width)
    
    set_table_borders(table2)
    
    # Add headers for second table
    header_row2 = table2.rows[0]
    for i, header in enumerate(headers2):
        cell = header_row2.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(cell)
    
    # Add data rows for second table
    for activity in activities_data:
        row = table2.add_row()
        
        # Evidence Admissibility
        row.cells[0].text = activity['evidence_admissibility']
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(row.cells[0])
        
        # Evidence Quality
        quality_cell = row.cells[1]
        quality_cell.text = activity['evidence_quality']
        quality_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in quality_cell.paragraphs:
            if "STRONG" in activity['evidence_quality']:
                paragraph.runs[0].font.color.rgb = RGBColor(0, 100, 0)
            elif "WEAK" in activity['evidence_quality']:
                paragraph.runs[0].font.color.rgb = RGBColor(156, 87, 0)
            elif "INADMISSIBLE" in activity['evidence_quality']:
                paragraph.runs[0].font.color.rgb = RGBColor(156, 0, 6)
        set_cell_borders(quality_cell)
        
        # Design Effectiveness
        row.cells[2].text = activity['required_effectiveness_design']
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(row.cells[2])
        
        # Implementation Effectiveness
        row.cells[3].text = activity['required_effectiveness_implementation']
        row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(row.cells[3])
        
        # Operating Effectiveness
        row.cells[4].text = activity['required_effectiveness_operating']
        row.cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(row.cells[4])
    
    document.add_paragraph()  # Add spacing after tables

def get_clause_severity(clause_id):
    """
    Get the overall severity for a clause (highest severity across all applicable activities)
    This matches the logic in your activity route
    """
    try:
        # Get all control activities for this clause through the proper join path
        control_activities = (
            db.session.query(ProjectControlActivity)
            .join(
                ProjectComplianceActivity,
                ProjectControlActivity.project_compliance_activity_id == ProjectComplianceActivity.id
            )
            .filter(ProjectComplianceActivity.project_clause_id == clause_id)
            .filter(ProjectComplianceActivity.applicability == True)  # Only applicable activities
            .all()
        )
        
        # If there are no applicable activities, return "No Activities"
        if not control_activities:
            return "No Activities"
        
        # Determine highest severity
        severity_hierarchy = {
            'Critical': 5,
            'Major': 4,
            'Significant': 3,
            'Minor': 2,
            'No findings noted': 1,
            'Not Classified': 0
        }
        
        highest_severity = 'No findings noted'
        highest_score = 0
        
        for activity in control_activities:
            activity_severity = activity.overall_severity_classification
            
            # Default to 'No findings noted' if nothing found and activity is compliant
            if not activity_severity or activity_severity == 'Not Classified':
                if activity.compliant_status == 'Compliant':
                    activity_severity = 'No findings noted'
                else:
                    activity_severity = 'Not Classified'
            
            if activity_severity and activity_severity != 'Not Classified':
                severity_score = severity_hierarchy.get(activity_severity, 0)
                if severity_score > highest_score:
                    highest_score = severity_score
                    highest_severity = activity_severity
        
        return highest_severity
        
    except Exception as e:
        current_app.logger.error(f"Error getting severity for clause {clause_id}: {e}")
        return "Unknown"